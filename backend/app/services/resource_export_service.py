from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from uuid import uuid4

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models import LearningResource, ReviewReport


EXPORT_ROOT = Path(__file__).resolve().parents[3] / "storage" / "exports"

# reportlab CID fonts are registered lazily (they pull in heavy PDF machinery).
_FONTS_REGISTERED = False


def _safe_stem(value: str) -> str:
    normalized = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', " ", value)
    normalized = " ".join(normalized.split()).strip(". ")
    return normalized[:80]


def _export_file_name(resource: LearningResource, audience: str, suffix: str) -> str:
    file_stem = _safe_stem(resource.title) or _safe_stem(resource.public_id) or "learning-resource"
    audience_label = {"learner": "学习者版", "teacher": "教师版"}[audience]
    return f"{file_stem}_v{resource.version}_{audience_label}{suffix}"


def _export_content(resource: LearningResource, audience: str) -> str:
    if resource.resource_type != "graded_quiz" or audience == "teacher":
        return resource.content_md
    hidden_prefixes = ("参考答案：", "答案：", "解析：")
    return "\n".join(
        line for line in resource.content_md.splitlines()
        if not line.strip().startswith(hidden_prefixes)
    )


# ---------------------------------------------------------------------------
# Markdown -> blocks
# ---------------------------------------------------------------------------


@dataclass
class _Block:
    kind: str
    level: int = 0
    text: str = ""
    items: list[str] = field(default_factory=list)
    ordered: bool = False
    code: str = ""
    header: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


_HR_RE = re.compile(r"^(\*\*\*|---|___)\s*$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _parse_blocks(markdown: str) -> list[_Block]:
    lines = markdown.split("\n")
    blocks: list[_Block] = []
    i = 0
    n = len(lines)
    while i < n:
        stripped = lines[i].strip()

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            blocks.append(_Block(kind="code", code="\n".join(code_lines).rstrip("\n")))
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            blocks.append(
                _Block(kind="heading", level=len(heading.group(1)), text=heading.group(2).strip())
            )
            i += 1
            continue

        if _HR_RE.match(stripped):
            blocks.append(_Block(kind="hr"))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if table_lines:
                header = _split_table_row(table_lines[0])
                rows: list[list[str]] = []
                for row_line in table_lines[1:]:
                    if re.fullmatch(r"[\s|:\-]+", row_line):
                        continue  # alignment separator
                    rows.append(_split_table_row(row_line))
                blocks.append(_Block(kind="table", header=header, rows=rows))
            continue

        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append(_Block(kind="quote", text="\n".join(quote_lines)))
            continue

        bullet = _BULLET_RE.match(stripped)
        ordered = _ORDERED_RE.match(stripped)
        if bullet or ordered:
            items: list[str] = []
            is_ordered = ordered is not None
            while i < n:
                s = lines[i].strip()
                bm = _BULLET_RE.match(s)
                om = _ORDERED_RE.match(s)
                if is_ordered and om:
                    items.append(om.group(1).strip())
                    i += 1
                elif not is_ordered and bm:
                    items.append(bm.group(1).strip())
                    i += 1
                else:
                    break
            blocks.append(_Block(kind="list", items=items, ordered=is_ordered))
            continue

        if not stripped:
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < n:
            s = lines[i].strip()
            if not s:
                break
            if (
                _HEADING_RE.match(s)
                or s.startswith("```")
                or s.startswith("|")
                or s.startswith(">")
                or _BULLET_RE.match(s)
                or _ORDERED_RE.match(s)
                or _HR_RE.match(s)
            ):
                break
            para_lines.append(s)
            i += 1
        blocks.append(_Block(kind="paragraph", text="\n".join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------


@dataclass
class _Inline:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    link: str | None = None


_INLINE_RE = re.compile(r"\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\)")


def _parse_inline(text: str) -> list[_Inline]:
    segments: list[_Inline] = []
    pos = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > pos:
            segments.append(_Inline(text=text[pos : match.start()]))
        if match.group(2) is not None:
            segments.append(_Inline(text=match.group(2), bold=True))
        elif match.group(3) is not None:
            segments.append(_Inline(text=match.group(3), italic=True))
        elif match.group(4) is not None:
            segments.append(_Inline(text=match.group(4), code=True))
        elif match.group(5) is not None:
            segments.append(_Inline(text=match.group(5), link=match.group(6)))
        pos = match.end()
    if pos < len(text):
        segments.append(_Inline(text=text[pos:]))
    return segments or [_Inline(text=text)]


def _escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _is_ascii(value: str) -> bool:
    return all(ord(char) < 128 for char in value)


# ---------------------------------------------------------------------------
# PDF rendering (reportlab)
# ---------------------------------------------------------------------------


def _inline_pdf(segments: list[_Inline]) -> str:
    parts: list[str] = []
    for seg in segments:
        text = _escape(seg.text)
        if seg.code:
            face = "Courier" if _is_ascii(seg.text) else "STSong-Light"
            text = f'<font face="{face}" color="#b03a2e">{text}</font>'
        if seg.link:
            text = f'<a href="{_escape(seg.link)}" color="#1a6fb5">{text}</a>'
        if seg.bold:
            text = f"<b>{text}</b>"
        if seg.italic:
            text = f"<i>{text}</i>"
        parts.append(text)
    return "".join(parts)


def _register_pdf_fonts() -> None:
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    pdfmetrics.registerFontFamily(
        "STSong-Light",
        normal="STSong-Light",
        bold="STSong-Light",
        italic="STSong-Light",
        boldItalic="STSong-Light",
    )
    _FONTS_REGISTERED = True


def _write_pdf(path: Path, resource: LearningResource, content: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    _register_pdf_fonts()

    navy = colors.HexColor("#1f3a5f")
    slate = colors.HexColor("#3d5a80")
    muted = colors.HexColor("#6b7280")
    code_bg = colors.HexColor("#f5f6f8")
    code_border = colors.HexColor("#d9dde3")

    title_style = ParagraphStyle(
        "ExportTitle",
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        textColor=navy,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "ExportH2",
        fontName="STSong-Light",
        fontSize=14,
        leading=20,
        textColor=navy,
        spaceBefore=16,
        spaceAfter=6,
    )
    h3_style = ParagraphStyle(
        "ExportH3",
        fontName="STSong-Light",
        fontSize=12,
        leading=17,
        textColor=slate,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ExportBody",
        fontName="STSong-Light",
        fontSize=10.5,
        leading=17,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor("#1f2937"),
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "ExportBullet",
        parent=body_style,
        leftIndent=16,
        firstLineIndent=-12,
        spaceAfter=3,
    )
    ordered_style = ParagraphStyle(
        "ExportOrdered",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=-14,
        spaceAfter=3,
    )
    quote_style = ParagraphStyle(
        "ExportQuote",
        parent=body_style,
        leftIndent=14,
        textColor=muted,
    )
    code_style = ParagraphStyle(
        "ExportCode",
        fontName="Courier",
        fontSize=8.5,
        leading=12.5,
        textColor=colors.HexColor("#24292f"),
    )
    cell_style = ParagraphStyle(
        "ExportCell",
        fontName="STSong-Light",
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#1f2937"),
    )
    cell_head_style = ParagraphStyle(
        "ExportCellHead",
        parent=cell_style,
        textColor=navy,
    )

    def footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(muted)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, f"第 {doc.page} 页")
        canvas.restoreState()

    story: list[object] = []
    for block in _parse_blocks(content):
        if block.kind == "heading":
            if block.level == 1:
                story.append(Paragraph(_inline_pdf(_parse_inline(block.text)), title_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1.2,
                        color=navy,
                        spaceBefore=2,
                        spaceAfter=12,
                    )
                )
            elif block.level == 2:
                story.append(Paragraph(_inline_pdf(_parse_inline(block.text)), h2_style))
            else:
                story.append(Paragraph(_inline_pdf(_parse_inline(block.text)), h3_style))
        elif block.kind == "paragraph":
            story.append(Paragraph(_inline_pdf(_parse_inline(block.text)), body_style))
        elif block.kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color=code_border, spaceBefore=6, spaceAfter=6))
        elif block.kind == "quote":
            story.append(Paragraph(_inline_pdf(_parse_inline(block.text)), quote_style))
        elif block.kind == "list":
            for index, item in enumerate(block.items, start=1):
                if block.ordered:
                    story.append(Paragraph(f"{index}. {_inline_pdf(_parse_inline(item))}", ordered_style))
                else:
                    story.append(Paragraph(f"• {_inline_pdf(_parse_inline(item))}", bullet_style))
        elif block.kind == "code":
            face = "Courier" if _is_ascii(block.code) else "STSong-Light"
            code_markup = "<br/>".join(_escape(line) for line in block.code.split("\n"))
            wrapped = Paragraph(f'<font face="{face}">{code_markup}</font>', code_style)
            code_table = Table([[wrapped]], colWidths=[A4[0] - 4 * cm])
            code_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), code_bg),
                        ("BOX", (0, 0), (-1, -1), 0.5, code_border),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ]
                )
            )
            story.append(code_table)
            story.append(Spacer(1, 8))
        elif block.kind == "table":
            data = [[Paragraph(_inline_pdf(_parse_inline(cell)), cell_head_style) for cell in block.header]]
            for row in block.rows:
                data.append([Paragraph(_inline_pdf(_parse_inline(cell)), cell_style) for cell in row])
            width = A4[0] - 4 * cm
            col_widths = [width / max(1, len(block.header))] * max(1, len(block.header))
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3f8")),
                        ("GRID", (0, 0), (-1, -1), 0.4, code_border),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        title=resource.title,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2.2 * cm,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


# ---------------------------------------------------------------------------
# Word rendering (python-docx)
# ---------------------------------------------------------------------------


def _docx_set_run(run, *, size: float, mono: bool = False, bold: bool | None = None, color: str | None = None) -> None:
    from docx.shared import Pt, RGBColor

    run.font.name = "Consolas" if mono else "Calibri"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")


def _docx_shade(paragraph, fill: str = "F2F3F5") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _write_docx(path: Path, resource: LearningResource, content: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )

    for block in _parse_blocks(content):
        if block.kind == "heading":
            if block.level == 1:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(block.text)
                _docx_set_run(run, size=20, bold=True, color="1F3A5F")
            elif block.level == 2:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block.text)
                _docx_set_run(run, size=14, bold=True, color="1F3A5F")
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block.text)
                _docx_set_run(run, size=12, bold=True, color="3D5A80")
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph()
            for seg in _parse_inline(block.text):
                run = paragraph.add_run(seg.text)
                _docx_set_run(
                    run,
                    size=10.5,
                    mono=seg.code,
                    bold=seg.bold,
                    color="B03A2E" if seg.code else "1F2937",
                )
        elif block.kind == "quote":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(16)
            for seg in _parse_inline(block.text):
                run = paragraph.add_run(seg.text)
                _docx_set_run(run, size=10.5, color="6B7280")
        elif block.kind == "hr":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run("─" * 40)
            _docx_set_run(run, size=8, color="D9DDE3")
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = doc.add_paragraph(style=style)
                for seg in _parse_inline(item):
                    run = paragraph.add_run(seg.text)
                    _docx_set_run(
                        run,
                        size=10.5,
                        mono=seg.code,
                        bold=seg.bold,
                        color="B03A2E" if seg.code else "1F2937",
                    )
        elif block.kind == "code":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(12)
            _docx_shade(paragraph, "F2F3F5")
            for line_number, line in enumerate(block.code.split("\n")):
                if line_number:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(line)
                _docx_set_run(run, size=9, mono=True, color="24292F")
        elif block.kind == "table":
            table = doc.add_table(rows=1, cols=max(1, len(block.header)))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for index, cell in enumerate(block.header):
                header_cells[index].text = ""
                run = header_cells[index].paragraphs[0].add_run(cell)
                _docx_set_run(run, size=9, bold=True, color="1F3A5F")
            for row in block.rows:
                cells = table.add_row().cells
                for index, cell in enumerate(row):
                    if index >= len(cells):
                        break
                    cells[index].text = ""
                    run = cells[index].paragraphs[0].add_run(cell)
                    _docx_set_run(run, size=9)

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Export entry point
# ---------------------------------------------------------------------------


def export_resource(
    db: Session,
    resource: LearningResource,
    export_format: str,
    audience: str = "learner",
) -> dict:
    export_format = export_format.lower()
    if export_format not in {"markdown", "pdf", "word"}:
        raise ValueError("export_format must be markdown, pdf or word")
    if audience not in {"learner", "teacher"}:
        raise ValueError("audience must be learner or teacher")
    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    suffix = {"markdown": ".md", "pdf": ".pdf", "word": ".docx"}[export_format]
    export_id = f"exp_{uuid4().hex}"
    path = EXPORT_ROOT / _export_file_name(resource, audience, suffix)
    content = _export_content(resource, audience)
    if export_format == "markdown":
        path.write_text(
            f"# {resource.title}\n\n{content}\n\n"
            f"---\n资源版本：{resource.version}\n审核状态：{resource.review_status}\n",
            encoding="utf-8",
        )
    elif export_format == "pdf":
        _write_pdf(path, resource, content)
    else:
        _write_docx(path, resource, content)
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    report = db.scalar(
        select(ReviewReport)
        .where(ReviewReport.resource_id == resource.id)
        .order_by(ReviewReport.id.desc())
    )
    return {
        "export_id": export_id,
        "resource_id": resource.public_id,
        "resource_version": resource.version,
        "format": export_format,
        "audience": audience,
        "file_name": path.name,
        "file_hash": f"sha256:{digest}",
        "review_report_id": str(report.id) if report else None,
        "review_status": resource.review_status,
        "download_url": f"/api/v1/resources/exports/{path.name}",
    }


def resolve_export_path(file_name: str) -> Path:
    candidate = (EXPORT_ROOT / Path(file_name).name).resolve()
    if candidate.parent != EXPORT_ROOT.resolve() or not candidate.is_file():
        raise FileNotFoundError(file_name)
    return candidate
